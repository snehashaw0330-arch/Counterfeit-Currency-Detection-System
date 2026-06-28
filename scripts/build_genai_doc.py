"""Generate a plain-language Word (.docx) explainer of the project's AI features.

Produces docs/GENAI_Explained.docx — a shareable, non-technical document
describing the Generative-AI features, the un-clonable circular security
pattern, and the various machine-learning techniques. Run:

    python scripts/build_genai_doc.py

Reproducible: re-running overwrites the same file. Needs python-docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "docs" / "GENAI_Explained.docx"

ACCENT = RGBColor(0x1F, 0x6F, 0x6B)   # teal — matches the app's "security ink"
MUTED = RGBColor(0x66, 0x66, 0x66)


def add_title(doc, text, subtitle):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = MUTED


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Pt(36 * level)
    return p


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def quote(doc, text):
    p = doc.add_paragraph(style="Intense Quote")
    p.add_run(text)
    return p


def rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("• • •")
    r.font.color.rgb = MUTED


def build():
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_title(
        doc,
        "The AI Side of Our Project — Explained Simply",
        "Counterfeit Currency Detection System · built by Sneha Shaw",
    )

    para(
        doc,
        "This note explains, in plain English, the “smart” parts of our "
        "project — the Generative AI features, the un-clonable circular "
        "security pattern, and the various machine-learning techniques we use to "
        "decide if a note is real or fake. No technical background needed.",
    )

    rule(doc)

    # ---- one-line summary ----
    h1(doc, "The one-line summary")
    quote(
        doc,
        "We built an app where you take a photo of an Indian banknote and it tells "
        "you if it looks REAL, SUSPICIOUS, FAKE, or “can’t tell — "
        "retake the photo.” On top of that it has three AI helpers: one that "
        "explains the result in simple words, one that draws an un-copyable "
        "security pattern, and a chat assistant that answers your questions about "
        "the app.",
    )
    para(doc, "Everything runs on your own computer. No internet needed, nothing is uploaded.")

    rule(doc)

    # ---- Part 1: the brain ----
    h1(doc, "Part 1 — The brain: how it decides real vs fake (the “many techniques” bit)")
    para(
        doc,
        "Instead of trusting one opinion, the app asks several different experts and "
        "combines their answers — exactly why the project is called "
        "“…with Various Machine Learning Techniques.”",
    )
    para(doc, "Think of it like a panel of judges, each looking at the note differently:")

    numbered(
        doc,
        "A deep-learning “eye” (a neural network called MobileNetV2). It "
        "looks at the whole note like a human glancing at it and gives a gut "
        "feeling: real or fake. Fast, but it can be fooled by unusual photos — "
        "so we don’t let it decide alone.",
    )
    numbered(
        doc,
        "A panel of about 15 forensic checks. These are small, focused tests, each "
        "checking one security feature the way a bank teller would — the serial "
        "number, the denomination, the size and shape vs the real RBI measurements, "
        "the watermark, security thread, Gandhi portrait, colour, micro-printing, "
        "bleed lines, the touch-mark, and more. Each test honestly says PASS, FAIL, "
        "or “couldn’t tell” — it never fakes a result.",
    )
    numbered(
        doc,
        "A second team of classical ML models (with names like SVM, Random Forest, "
        "KNN). They study measurable clues — texture, colour, structure — "
        "and give an independent “second opinion.” On our data these "
        "actually beat the fancy neural network, which is why showing several "
        "techniques matters.",
    )

    para(
        doc,
        "The final verdict is all of these fused together, with a few “hard "
        "rules”: if the structure, colour, or size are clearly wrong, the note "
        "is failed no matter what — the way one glaring red flag overrides "
        "everything else.",
    )
    quote(
        doc,
        "Why this is good: one model can be confidently wrong. A panel that has to "
        "agree is much harder to fool, and we can show you exactly which check "
        "caused the verdict.",
    )

    rule(doc)

    # ---- Part 2: Explain with AI ----
    h1(doc, "Part 2 — Generative AI feature #1: “Explain with AI”")
    para(
        doc,
        "A raw result is full of jargon (“hue entropy”, “aspect "
        "deviation”…). Most people don’t want that. So we added a "
        "button — Explain with AI — that turns the technical output into a "
        "short, calm, plain-language explanation, like:",
    )
    quote(
        doc,
        "“This ₹500 note looks likely genuine. The watermark, security "
        "thread and colour all checked out. This isn’t a 100% guarantee — "
        "if it matters, also check it by hand: tilt it to see the number change "
        "colour…”",
    )
    para(doc, "What makes it nice:")
    bullet(
        doc,
        "It also tells you what to check by hand — tilt for the colour shift, "
        "hold to light for the watermark, feel the raised print. Useful in the real world.",
    )
    bullet(
        doc,
        "It can be read out loud and works in Hindi too — built with "
        "visually-impaired users in mind (the same people RBI’s touch-features are for).",
    )
    bullet(
        doc,
        "It always works. When an AI key is available it uses Claude (a large "
        "language model) to write the explanation. When there’s no key or no "
        "internet, it falls back to a built-in template that says the same thing. "
        "The user never sees an error.",
    )
    quote(
        doc,
        "Important and honest: this AI explains a result. It does NOT generate money "
        "or claim certainty. It openly says the app can be wrong on good fakes.",
    )

    rule(doc)

    # ---- Part 3: the circle ----
    h1(doc, "Part 3 — The “un-clonable circle”: Security Pattern Studio")
    para(doc, "This is the circular pattern feature you remember. Here’s the simple idea.")
    para(
        doc,
        "The problem with copying money isn’t the picture of Gandhi — "
        "it’s the incredibly fine, woven, swirling line-work printed on real "
        "notes. Those swirls are called guilloché. They look like the spiro-art "
        "you may have drawn as a kid with a gear and a pen, but far denser. They are "
        "easy for a computer to draw perfectly from a formula, but very hard to copy "
        "by hand or through a scan-and-print — a photocopier smudges the fine lines.",
    )
    para(doc, "Our Security Pattern Studio generates exactly this:")
    bullet(doc, "You give it a seed — any word or number (for example a serial number).")
    bullet(
        doc,
        "It instantly draws a unique, intricate circular pattern: woven rings on the "
        "outside, a tiny ring of micro-text, and a layered rosette (spirograph-style "
        "swirls) in the middle.",
    )
    bullet(
        doc,
        "Same seed → exactly the same pattern, every time, on any computer. A "
        "different seed → a completely different pattern.",
    )
    para(doc, "Why that’s the “un-clonable” part, in plain terms:")
    bullet(
        doc,
        "The pattern is mathematically exact and reproducible, so a genuine pattern "
        "for a given serial can always be re-generated and compared — if a copy "
        "doesn’t match precisely, it’s suspect.",
    )
    bullet(
        doc,
        "The fine, interwoven detail is what defeats casual copying — the same "
        "principle real security printers rely on.",
    )
    quote(
        doc,
        "What it is NOT (and this is deliberate): it does NOT create fake notes or "
        "anything that looks like currency — no Gandhi, no “RBI”, no "
        "rupee layout. It only makes abstract decorative art that demonstrates the "
        "anti-copy design idea. Making realistic fake-note images would be illegal, "
        "so we intentionally did not build that. This is the legitimate, honest "
        "version of the “make something that can’t be cloned” idea.",
    )

    # Embed the sample pattern image if it exists.
    sample = OUT.parent / "sample_security_pattern.png"
    if sample.exists():
        from docx.shared import Inches
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(sample), width=Inches(3.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run("An example generated security pattern (abstract art — not currency).")
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED

    rule(doc)

    # ---- Part 4: chatbot ----
    h1(doc, "Part 4 — Generative AI feature #2: the Help Chatbot")
    para(
        doc,
        "There’s a little chat assistant in the corner of the app. You can ask "
        "it anything about the project in normal language:",
    )
    bullet(doc, "“How does it work?”")
    bullet(doc, "“What does SUSPICIOUS mean?”")
    bullet(doc, "“How accurate is it?”")
    bullet(doc, "“What is the security pattern / the AI heatmap / Hindi mode?”")
    bullet(doc, "“How do I run it?”")
    para(doc, "How it behaves:")
    bullet(
        doc,
        "It’s friendly and knows this project inside-out — it’s loaded "
        "with a curated summary of how everything works.",
    )
    bullet(
        doc,
        "Like the explainer, it uses Claude when a key is available, and a built-in "
        "question-and-answer brain when offline — so it always replies, even "
        "with no internet, during a demo.",
    )
    bullet(
        doc,
        "It only talks about this project. If someone asks it how to make or pass off "
        "fake money, it politely refuses — it’s a verification tool, not a how-to.",
    )
    bullet(doc, "(Small fun touch: it knows Sneha built it.)")

    rule(doc)

    # ---- wrap-up ----
    h1(doc, "Putting it together — what to tell someone in 20 seconds")
    bullet(doc, "You photograph a note, the app gives an honest verdict and shows why.")
    bullet(
        doc,
        "It uses many techniques at once (a neural net + ~15 forensic checks + a "
        "classical-ML second opinion) instead of trusting one — harder to fool.",
    )
    bullet(doc, "Three AI helpers sit on top:")
    bullet(
        doc,
        "Explain with AI — plain-language result + hand-checks, in English/Hindi, "
        "can be read aloud.",
        level=1,
    )
    bullet(
        doc,
        "Security Pattern Studio — generates the un-clonable circular guilloché "
        "pattern (abstract art, not currency).",
        level=1,
    )
    bullet(doc, "Help Chatbot — answers questions about the app, refuses misuse.", level=1)
    bullet(
        doc,
        "It’s honest about limits: it’s a screening aid on a phone photo, not "
        "a guarantee, and it never pretends to be certain.",
    )

    rule(doc)
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = closing.add_run("Questions? Just ask the in-app chat assistant — or ping Sneha.")
    cr.italic = True
    cr.font.color.rgb = MUTED

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
